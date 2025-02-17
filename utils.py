import os
from openai import OpenAI
import docx
import PyPDF2

def load_prompt(prompt_path: str) -> str:
    """Reads the prompt file and returns its content as a string."""
    with open(prompt_path, 'r', encoding='utf-8') as f:
        return f.read()

def load_news(input_folder: str) -> list:
    """
    Returns a list of tuples (file_name, text_content).
    Supports .txt and .pdf files inside the input_folder.
    """
    files = []
    for file_name in os.listdir(input_folder):
        file_path = os.path.join(input_folder, file_name)
        if file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            files.append((file_name, content))
        elif file_name.endswith('.pdf'):
            content = extract_text_from_pdf(file_path)
            files.append((file_name, content))
    return files

def load_news_from_list(docs_list: list) -> list:
    """
    Returns a list of tuples (file_name, text_content).
    Supports .txt and .pdf files inside the input_folder.
    """
    files = []
    for file_name in docs_list:
        file_path = file_name
        if file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            files.append((file_name, content))
        elif file_name.endswith('.pdf'):
            content = extract_text_from_pdf(file_path)
            files.append((file_name, content))
    return files

def load_news_from_list(docs_list: list) -> list:
    """
    Returns a list of tuples (file_name, text_content).
    Supports .txt and .pdf files inside the input_folder.
    """
    files = []
    for file_name in docs_list:
        #file_path = os.path.join(input_folder, file_name)
        file_path = file_name
        if file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            files.append((file_name, content))
        elif file_name.endswith('.pdf'):
            content = extract_text_from_pdf(file_path)
            files.append((file_name, content))
    return files

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extracts text content from a PDF file."""
    text = ""
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text()
    return text

def generate_summary(client: OpenAI, text: str, prompt: str, model: str = "gpt-4") -> dict:
    """
    Sends the text to OpenAI's API to generate a summary in Portuguese
    according to the instructions in the prompt.

    Returns:
        dict: A dictionary with keys 'title' and 'content'.
    """
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": text}
    ]
    
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.7,
        max_tokens=1000
    )

    full_summary = response.choices[0].message.content.strip()
    
    lines = full_summary.split("\n", 1)
    title = lines[0].strip() if lines else "No Title"
    content = lines[1].strip() if len(lines) > 1 else "No Content"
    
    return {"title": title, "content": content}

def save_summary(output_folder: str, file_name: str, summary: dict, format: str = "docx"):
    """
    Saves the summary in the specified format (.docx or .txt) inside the 'output_folder'.

    Args:
        output_folder (str): Path to the output folder.
        file_name (str): Name of the output file.
        summary (dict): Dictionary with 'title' and 'content'.
        format (str): Output format, either 'docx' or 'txt'.
    """
    if format == "docx":
        doc = docx.Document()
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(summary["title"])
        title_run.bold = True
        doc.add_paragraph(summary["content"])
        doc.add_paragraph("________________________________________")
        output_path = os.path.join(output_folder, file_name + ".docx")
        doc.save(output_path)
    elif format == "txt":
        output_path = os.path.join(output_folder, file_name + ".txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"{summary['title']}\n\n{summary['content']}\n")
    else:
        raise ValueError("Unsupported format. Choose 'docx' or 'txt'.")